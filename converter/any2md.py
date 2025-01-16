import logging
import os
from io import BytesIO
import pdfplumber
from ebooklib import epub
from bs4 import BeautifulSoup
import pandas as pd
from markdownify import markdownify
from odf.opendocument import load as load_odf
from odf.text import P, H
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def convert_pdf_to_markdown(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages])
        return markdownify(text)
    except Exception as e:
        logging.error(f"Failed to convert PDF to Markdown: {e}")
        return ""

def convert_epub_to_markdown(file_path):
    try:
        book = epub.read_epub(file_path)
        content = []
        for item in book.get_items():
            if item.get_type() == 9:  # EPUB.TEXT
                soup = BeautifulSoup(item.get_content(), "html.parser")
                content.append(markdownify(str(soup)))
        return "\n".join(content)
    except Exception as e:
        logging.error(f"Failed to convert EPUB to Markdown: {e}")
        return ""

def convert_html_to_markdown(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            html_content = file.read()
        return markdownify(html_content)
    except Exception as e:
        logging.error(f"Failed to convert HTML to Markdown: {e}")
        return ""

def convert_csv_to_markdown(file_path):
    try:
        df = pd.read_csv(file_path)
        return df.to_markdown(index=False)
    except Exception as e:
        logging.error(f"Failed to convert CSV to Markdown: {e}")
        return ""

def convert_txt_to_markdown(file_or_content):
    try:
        if os.path.isfile(file_or_content):
            with open(file_or_content, "r", encoding="utf-8") as file:
                return file.read()
        else:
            return file_or_content
    except Exception as e:
        logging.error(f"Failed to convert TXT to Markdown: {e}")
        return ""

def convert_docx_to_markdown(file_or_content):
    try:
        if isinstance(file_or_content, str):
            doc = Document(file_or_content)
        else:
            doc = Document(BytesIO(file_or_content))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logging.error(f"Failed to convert DOCX to Markdown: {e}")
        return ""

def convert_odt_to_markdown(file_path):
    try:
        if isinstance(file_path, str) and os.path.isfile(file_path):
            with open(file_path, "rb") as f:
                odt_content = f.read()
        elif isinstance(file_path, bytes):
            odt_content = file_path
        else:
            raise ValueError("Invalid ODT input. Provide a valid file path or bytes content.")

        odt_file = load_odf(BytesIO(odt_content))
        content = []
        for element in odt_file.getElementsByType(P) + odt_file.getElementsByType(H):
            text = element.getText().strip()
            if text:
                if element.qname.localname == "H":
                    level = int(element.getAttribute("outline-level"))
                    content.append(f"{'#' * level} {text}")
                else:
                    content.append(text)
        return "\n\n".join(content)
    except Exception as e:
        logging.error(f"Failed to convert ODT to Markdown: {e}")
        return ""

def convert_to_markdown(content_or_path, ext):
    try:
        if os.path.isfile(content_or_path):
            content = content_or_path
        else:
            content = BytesIO(content_or_path)

        if ext == ".pdf":
            return convert_pdf_to_markdown(content)
        elif ext == ".epub":
            return convert_epub_to_markdown(content)
        elif ext in [".html", ".htm"]:
            return convert_html_to_markdown(content)
        elif ext == ".csv":
            return convert_csv_to_markdown(content)
        elif ext == ".txt":
            return convert_txt_to_markdown(content)
        elif ext == ".docx":
            return convert_docx_to_markdown(content)
        elif ext == ".odt":
            return convert_odt_to_markdown(content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        logging.error(f"Failed to convert file to Markdown: {e}")
        return ""